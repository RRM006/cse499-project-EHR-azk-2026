"""MEDIC-6/7 — post-referral summary data + working .docx download, fully offline.

Must: expose the patient (with vitals) on GET /visits/{uuid}; let staff edit
weight/BP (audit-logged, guarded); render the C1 suggested condition WITH its
embedded disclaimer in the summary_report .docx; and regenerate the report fresh
at download time so staff edits made after an earlier report still show (the
staleness half of "download must actually work").
"""

from io import BytesIO

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from backend.app.db.database import Base, get_db
from backend.app.db.models import AuditLog, CaseProfile, Clinic, Patient, User, Utterance, Visit
from backend.app.main import app
from backend.app.services.documents.storage import FilesystemStorage
from backend.app.services.suggestion import CONDITION_DISCLAIMER


@pytest.fixture()
def env(tmp_path, monkeypatch):
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False},
                           poolclass=StaticPool, future=True)
    Base.metadata.create_all(bind=engine)
    TestSession = sessionmaker(bind=engine, autoflush=False, future=True)

    db = TestSession()
    clinic = Clinic(name="Demo Clinic")
    db.add(clinic)
    db.flush()
    patient = Patient(clinic_id=clinic.id, external_ref="+8801712345678",
                      display_name="Kamal Hossain", birth_year=1985, sex="male")
    medic = User(clinic_id=clinic.id, name="Medic Rahman", role="medic")
    desk = User(clinic_id=clinic.id, name="Desk Clerk", role="desk")
    db.add_all([patient, medic, desk])
    db.flush()
    visit = Visit(clinic_id=clinic.id, patient_id=patient.id, status="awaiting_doctor")
    db.add(visit)
    db.flush()
    db.add(Utterance(visit_id=visit.id, role="patient", seq=0,
                     raw_text="খাওয়ার পরে বুক জ্বালাপোড়া করে", source="mic"))
    db.add(CaseProfile(
        visit_id=visit.id,
        summary="Chest burning after meals.",
        entities={
            "summary_fields": {
                "main_problem": {"value": "Chest burning", "value_en": "Chest burning",
                                 "value_bn": "বুক জ্বালা", "source": "ai"},
            },
            "suggested_condition": {
                "condition": "GERD (Acid Reflux)", "condition_en": "GERD (Acid Reflux)",
                "condition_bn": "জিইআরডি", "reasoning_en": "Burning after meals fits reflux.",
                "reasoning_bn": "", "source": "ai", "edited_by": None, "edited_at": None,
                "disclaimer": CONDITION_DISCLAIMER, "disclaimer_bn": "রোগনির্ণয় নয়",
            },
        },
    ))
    db.commit()
    ids = {"visit_uuid": visit.uuid, "patient_id": patient.id,
           "medic_id": medic.id, "desk_id": desk.id}
    db.close()

    def override_get_db():
        s = TestSession()
        try:
            yield s
        finally:
            s.close()

    app.dependency_overrides[get_db] = override_get_db

    test_storage = FilesystemStorage(tmp_path)
    monkeypatch.setattr(
        "backend.app.services.documents.build_storage", lambda *a, **k: test_storage
    )
    monkeypatch.setattr(
        "backend.app.api.routes_documents.build_storage", lambda *a, **k: test_storage
    )

    yield TestClient(app), TestSession, ids
    app.dependency_overrides.clear()


def _docx_text(client, visit_uuid) -> str:
    r = client.post(f"/api/visits/{visit_uuid}/documents/summary_report")
    assert r.status_code == 200, r.text
    d = client.get(r.json()["download_url"])
    assert d.status_code == 200
    doc = DocxDocument(BytesIO(d.content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_visit_detail_embeds_patient_with_vitals(env):
    client, _, ids = env
    detail = client.get(f"/api/visits/{ids['visit_uuid']}").json()
    p = detail["patient"]
    assert p["display_name"] == "Kamal Hossain"
    assert p["external_ref"] == "+8801712345678"
    assert p["birth_year"] == 1985
    assert p["weight_kg"] is None and p["bp"] is None  # not recorded yet


def test_vitals_patch_updates_and_audits(env):
    client, TestSession, ids = env
    r = client.patch(f"/api/patients/{ids['patient_id']}/vitals",
                     json={"weight_kg": 72.5, "bp": "130/85", "editor_id": ids["medic_id"]})
    assert r.status_code == 200
    body = r.json()
    assert body["weight_kg"] == 72.5 and body["bp"] == "130/85"

    # Visible on the visit detail immediately (the medic screen re-reads this).
    assert client.get(f"/api/visits/{ids['visit_uuid']}").json()["patient"]["weight_kg"] == 72.5

    db = TestSession()
    log = db.query(AuditLog).filter(AuditLog.action == "patient.vitals_edit").one()
    assert log.actor_id == ids["medic_id"]
    assert log.detail == {"weight_kg": 72.5, "bp": "130/85"}
    db.close()


def test_vitals_guard_rails(env):
    client, _, ids = env
    pid, medic = ids["patient_id"], ids["medic_id"]

    # Non-staff editor -> 403.
    assert client.patch(f"/api/patients/{pid}/vitals",
                        json={"weight_kg": 70, "editor_id": ids["desk_id"]}).status_code == 403
    # Unknown patient -> 404.
    assert client.patch("/api/patients/99999/vitals",
                        json={"weight_kg": 70, "editor_id": medic}).status_code == 404
    # Invalid weight -> 422 (schema-validated).
    assert client.patch(f"/api/patients/{pid}/vitals",
                        json={"weight_kg": -3, "editor_id": medic}).status_code == 422
    assert client.patch(f"/api/patients/{pid}/vitals",
                        json={"weight_kg": 700, "editor_id": medic}).status_code == 422
    # Nothing to update -> 400.
    assert client.patch(f"/api/patients/{pid}/vitals",
                        json={"editor_id": medic}).status_code == 400


# --- S38 (A4/A5, rev 0013): height rides on the SAME row and the SAME endpoint ---


def test_height_is_recorded_through_the_existing_vitals_endpoint(env):
    """ADR-0060: height is the other half of a BMI and belongs with the vitals that are
    already on ``patients`` — not in a second store and not on a second route."""
    client, TestSession, ids = env
    r = client.patch(f"/api/patients/{ids['patient_id']}/vitals",
                     json={"height_cm": 168.5, "editor_id": ids["medic_id"]})
    assert r.status_code == 200, r.text
    assert r.json()["height_cm"] == 168.5
    # And it reaches the staff detail screens through the call they already make.
    assert client.get(f"/api/visits/{ids['visit_uuid']}").json()["patient"]["height_cm"] == 168.5

    db = TestSession()
    log = db.query(AuditLog).filter(AuditLog.action == "patient.vitals_edit").one()
    assert log.detail == {"height_cm": 168.5}   # only what was actually edited
    db.close()


def test_height_alone_is_enough_to_be_a_real_update(env):
    """The 'nothing to update' 400 must not fire when height is the only field sent —
    recording a height before a weight is a normal order of work."""
    client, _, ids = env
    assert client.patch(f"/api/patients/{ids['patient_id']}/vitals",
                        json={"height_cm": 150, "editor_id": ids["medic_id"]}).status_code == 200


def test_an_implausible_height_is_refused_by_the_schema(env):
    """Bounds mirror services/clinical_reference: a height the BMI calculator would
    refuse must not be storable, or the record keeps a value that silently shows no
    BMI with no explanation."""
    client, _, ids = env
    pid, medic = ids["patient_id"], ids["medic_id"]
    for bad in (17, 1.7, 300, -5):
        assert client.patch(f"/api/patients/{pid}/vitals",
                            json={"height_cm": bad, "editor_id": medic}).status_code == 422, bad


def test_no_bmi_is_ever_accepted_or_returned(env):
    """ADR-0060: BMI is derived on demand. If it were storable it would go stale the
    moment a weight was corrected, and two numbers would disagree in the record."""
    client, _, ids = env
    detail = client.get(f"/api/visits/{ids['visit_uuid']}").json()["patient"]
    assert not [k for k in detail if "bmi" in k.lower()]
    # An unknown field is ignored rather than stored (pydantic default), and certainly
    # never echoed back as if it had been accepted.
    r = client.patch(f"/api/patients/{ids['patient_id']}/vitals",
                     json={"height_cm": 170, "bmi": 99, "editor_id": ids["medic_id"]})
    assert r.status_code == 200
    assert "bmi" not in r.json()


def test_summary_docx_carries_condition_disclaimer_and_vitals(env):
    client, _, ids = env
    client.patch(f"/api/patients/{ids['patient_id']}/vitals",
                 json={"weight_kg": 72.5, "bp": "130/85", "editor_id": ids["medic_id"]})
    text = _docx_text(client, ids["visit_uuid"])

    # C1 section: label + condition + reasoning + the embedded disclaimer (rule #2).
    assert "Possible Condition (AI Suggestion – Not a Diagnosis)" in text
    assert "GERD (Acid Reflux)" in text and "জিইআরডি" in text
    assert "Burning after meals fits reflux." in text
    assert CONDITION_DISCLAIMER in text
    # Vitals in the patient meta; the M12 no-diagnosis disclaimer still present.
    assert "72.5 kg" in text and "130/85" in text
    assert "NOT a diagnosis" in text


def test_summary_docx_is_fresh_after_staff_edit(env):
    client, _, ids = env
    uuid = ids["visit_uuid"]
    # A report already exists from an earlier download...
    assert "Chest burning" in _docx_text(client, uuid)
    # ...then the medic corrects a field. The next download must show the new value.
    r = client.patch(f"/api/visits/{uuid}/profile/fields/main_problem",
                     json={"value": "Severe chest burning at night",
                           "editor_id": ids["medic_id"]})
    assert r.status_code == 200
    assert "Severe chest burning at night" in _docx_text(client, uuid)
