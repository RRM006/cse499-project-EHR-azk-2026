"""DOCTOR-6 — prescription Submit: persist the row, render the .docx, download it.

Guarantees checked: a prescription row + a linked document (kind 'prescription',
visit-grain) are created; the .docx renders the letterhead/patient/medicines and the
Diagnosis EXACTLY as typed; because the writer only reads the submitted payload, an
empty diagnosis stays empty — the AI suggested condition can never leak in (rule #2).
"""

from io import BytesIO

from datetime import date, timedelta

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from backend.app.db.database import Base, get_db
from backend.app.db.models import (
    CaseProfile, Clinic, Document, Patient, Prescription, User, Visit,
)
from backend.app.main import app
from backend.app.services.clinical_dates import dhaka_today_iso
from backend.app.services.documents.storage import FilesystemStorage


@pytest.fixture()
def env(tmp_path, monkeypatch):
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False},
                           poolclass=StaticPool, future=True)
    Base.metadata.create_all(bind=engine)
    TestSession = sessionmaker(bind=engine, autoflush=False, future=True)

    db = TestSession()
    clinic = Clinic(name="Demo Clinic", address="Dhanmondi, Dhaka")
    db.add(clinic)
    db.flush()
    patient = Patient(clinic_id=clinic.id, external_ref="+8801712345678",
                      display_name="Kamal Hossain", birth_year=1985, sex="male")
    doctor = User(clinic_id=clinic.id, name="Dr. M. Rahman", role="doctor",
                  qualification="MBBS, FCPS (Medicine)")
    medic = User(clinic_id=clinic.id, name="Medic Rahman", role="medic")
    db.add_all([patient, doctor, medic])
    db.flush()
    visit = Visit(clinic_id=clinic.id, patient_id=patient.id, status="awaiting_doctor")
    db.add(visit)
    db.flush()
    # A stored AI suggestion — the prescription must NEVER pull it into Diagnosis.
    db.add(CaseProfile(visit_id=visit.id, summary="x", entities={
        "suggested_condition": {"condition_en": "GERD (Acid Reflux)", "source": "ai"}}))
    db.commit()
    ids = {"visit_uuid": visit.uuid, "doctor_id": doctor.id, "medic_id": medic.id}
    db.close()

    def override_get_db():
        s = TestSession()
        try:
            yield s
        finally:
            s.close()

    app.dependency_overrides[get_db] = override_get_db
    test_storage = FilesystemStorage(tmp_path)
    monkeypatch.setattr(
        "backend.app.services.documents.build_storage", lambda *a, **k: test_storage
    )
    monkeypatch.setattr(
        "backend.app.api.routes_documents.build_storage", lambda *a, **k: test_storage
    )
    yield TestClient(app), TestSession, ids
    app.dependency_overrides.clear()


def _payload(diagnosis="Viral fever"):
    return {
        # S38 (B5): the prescription date must be TODAY in Dhaka and the follow-up must
        # not be in the past — both enforced server-side now. These were fixed literals
        # ("2026-07-06" / "2026-07-13"), which no assertion in this file ever looked at;
        # they are computed so the fixture stays valid on every future run rather than
        # rotting into a 400 tomorrow.
        "language": "en", "date": dhaka_today_iso(),
        "clinic": {"name": "Demo Clinic", "address": "Dhanmondi, Dhaka"},
        "doctor": {"id": 1, "name": "Dr. M. Rahman", "qualification": "MBBS, FCPS (Medicine)",
                   "registration_no": "BMDC-A-40001", "specialization": "Internal Medicine"},
        "patient": {"name": "Kamal Hossain", "phone": "+8801712345678", "age": "41",
                    "sex": "male", "weight_kg": "80", "bp": "120/80"},
        "symptoms": "Fever and cough for 3 days",
        "diagnosis": diagnosis,
        "medicines": [{"name": "Napa", "strength": "500mg", "dosage": "1+0+1",
                       "timing": "after meals", "duration": "5 days"}],
        "advice": "Rest and fluids", "tests": "CBC",
        "followup_date": (date.fromisoformat(dhaka_today_iso()) + timedelta(days=7)).isoformat(),
    }


def _docx_text(client, url) -> str:
    d = client.get(url)
    assert d.status_code == 200
    doc = DocxDocument(BytesIO(d.content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_submit_persists_row_and_links_document(env):
    client, TestSession, ids = env
    r = client.post(f"/api/visits/{ids['visit_uuid']}/prescription",
                    json={"doctor_id": ids["doctor_id"], "payload": _payload()})
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["prescription_id"] > 0
    assert body["document"]["kind"] == "prescription"
    assert body["document"]["download_url"].endswith("/download")

    with TestSession() as db:
        pres = db.get(Prescription, body["prescription_id"])
        assert pres is not None
        assert pres.doctor_id == ids["doctor_id"]
        assert pres.payload["diagnosis"] == "Viral fever"
        assert pres.document_id == body["document"]["id"]
        doc = db.get(Document, pres.document_id)
        assert doc.kind == "prescription"
        assert doc.utterance_id is None and doc.visit_id is not None  # visit-grain


def test_docx_renders_diagnosis_medicines_and_patient(env):
    client, _, ids = env
    r = client.post(f"/api/visits/{ids['visit_uuid']}/prescription",
                    json={"doctor_id": ids["doctor_id"], "payload": _payload()})
    text = _docx_text(client, r.json()["document"]["download_url"])
    assert "Demo Clinic" in text
    assert "Kamal Hossain" in text
    assert "Viral fever" in text          # diagnosis, exactly as typed
    assert "Napa" in text and "500mg" in text
    assert "CBC" in text


def test_empty_diagnosis_never_pulls_ai_condition(env):
    """The writer only reads the payload — a blank diagnosis stays blank and the
    stored AI suggestion (GERD) never appears in the prescription (rule #2)."""
    client, _, ids = env
    r = client.post(f"/api/visits/{ids['visit_uuid']}/prescription",
                    json={"doctor_id": ids["doctor_id"], "payload": _payload(diagnosis="")})
    text = _docx_text(client, r.json()["document"]["download_url"])
    assert "GERD" not in text and "Acid Reflux" not in text


def test_submit_non_doctor_400(env):
    client, _, ids = env
    r = client.post(f"/api/visits/{ids['visit_uuid']}/prescription",
                    json={"doctor_id": ids["medic_id"], "payload": _payload()})
    assert r.status_code == 400


def test_submit_unknown_visit_404(env):
    client, _, ids = env
    r = client.post("/api/visits/no-such-visit/prescription",
                    json={"doctor_id": ids["doctor_id"], "payload": _payload()})
    assert r.status_code == 404
