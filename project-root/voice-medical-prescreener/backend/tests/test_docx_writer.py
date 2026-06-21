"""Offline checks for the .docx export writer.

These never touch the network or a real DB: they build a writer over a plain
in-memory ``Utterance`` and inspect the produced bytes with python-docx. Raw and
corrected are now SEPARATE documents, so each is checked independently. The key
guard is rule #1 — the RAW text must appear in the raw document verbatim.
"""

from datetime import datetime, timezone
from io import BytesIO

import pytest
from docx import Document as DocxDocument

from backend.app.db.models import Utterance
from backend.app.services.documents import build_writer
from backend.app.services.documents.docx_writer import DocxWriter


def _sample_utterance() -> Utterance:
    utt = Utterance(
        raw_text="ami onek jor onuvob korchi   ",  # trailing spaces are part of raw
        corrected_text="আমি অনেক জ্বর অনুভব করছি",
        source="mic",
        stt_provider="browser_webspeech",
        correction_provider="gemini",
        correction_model="gemini-flash-latest",
    )
    utt.id = 12
    utt.created_at = datetime(2026, 6, 21, 9, 30, tzinfo=timezone.utc)
    utt.corrected_at = datetime(2026, 6, 21, 9, 31, tzinfo=timezone.utc)
    return utt


def _all_text(docx_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_build_writer_returns_docx_writer():
    writer = build_writer("docx")
    assert isinstance(writer, DocxWriter)
    assert writer.format == "docx"


def test_raw_document_holds_raw_verbatim_and_not_the_correction():
    utt = _sample_utterance()
    text = _all_text(DocxWriter().render(utt, kind="raw"))

    assert utt.raw_text in text  # RAW verbatim (rule #1), spaces and all
    assert str(utt.id) in text  # session id in the metadata
    assert "browser_webspeech" in text
    assert "Corrected Transcript" not in text  # raw doc is titled just "Transcript"
    assert utt.corrected_text not in text  # the correction is a SEPARATE file


def test_corrected_document_holds_the_correction_only():
    utt = _sample_utterance()
    text = _all_text(DocxWriter().render(utt, kind="corrected"))

    assert "Corrected Transcript" in text
    assert utt.corrected_text in text
    assert utt.raw_text not in text  # raw lives in its own file


def test_render_rejects_unknown_kind():
    with pytest.raises(ValueError):
        DocxWriter().render(_sample_utterance(), kind="combined")


def test_build_writer_rejects_unknown_format():
    with pytest.raises(ValueError):
        build_writer("xlsx")
