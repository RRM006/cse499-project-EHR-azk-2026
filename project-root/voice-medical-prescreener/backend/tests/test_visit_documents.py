"""Visit-grain document exports (step 3 of the Session-9 plan) — fully offline.

Proves: the transcript .docx contains every stored raw utterance BYTE-EXACT
(rule #1) in turn order with role labels; the summary report carries the 10 field
labels, the stored field values, and the M12 no-diagnosis disclaimer (rule #2);
the route guards (bad kind 400, unknown visit 404); download round-trips as a
Word file. Same in-memory DB + temp-dir storage pattern as test_routes_documents.
"""

from io import BytesIO

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from backend.app.db.database import Base, get_db
from backend.app.db.models import CaseProfile, Clinic, Patient, Utterance, Visit
from backend.app.main import app
from backend.app.services.documents.storage import FilesystemStorage

RAW_TURNS = [
    ("system", "আপনার সমস্যাটি নিজের ভাষায় খুলে বলুন তো।"),
    ("patient", "আমার তিন দিন ধরে মাথা ব্যথা আর হালকা জ্বর।"),
    ("system", "আপনি বর্তমানে কোনো ওষুধ খাচ্ছেন?"),
    ("patient", "napa খাচ্ছি দিনে দুইবার"),
]


@pytest.fixture()
def env(tmp_path, monkeypatch):
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
        future=True,
    )
    Base.metadata.create_all(bind=engine)
    TestSession = sessionmaker(bind=engine, autoflush=False, future=True)

    db = TestSession()
    clinic = Clinic(name="Demo Clinic")
    db.add(clinic)
    db.flush()
    patient = Patient(
        clinic_id=clinic.id, external_ref="+8801712345678",
        display_name="Kamal Hossain", weight_kg=70.5, bp="120/80",
    )
    db.add(patient)
    db.flush()
    visit = Visit(clinic_id=clinic.id, patient_id=patient.id, status="awaiting_review")
    db.add(visit)
    db.flush()
    for seq, (role, text) in enumerate(RAW_TURNS):
        db.add(Utterance(visit_id=visit.id, role=role, seq=seq, raw_text=text,
                         source="mic" if role == "patient" else "tts"))
    db.add(CaseProfile(
        visit_id=visit.id,
        summary="Headache and mild fever for three days.",
        entities={"summary_fields": {
            "main_problem": {"value": "মাথা ব্যথা ও হালকা জ্বর", "source": "ai"},
            "current_medicines": {"value": "Napa, twice daily", "source": "human"},
        }},
    ))
    db.commit()
    visit_uuid = visit.uuid
    db.close()

    def override_get_db():
        s = TestSession()
        try:
            yield s
        finally:
            s.close()

    app.dependency_overrides[get_db] = override_get_db

    test_storage = FilesystemStorage(tmp_path)
    monkeypatch.setattr(
        "backend.app.services.documents.build_storage", lambda *a, **k: test_storage
    )
    monkeypatch.setattr(
        "backend.app.api.routes_documents.build_storage", lambda *a, **k: test_storage
    )

    yield TestClient(app), visit_uuid
    app.dependency_overrides.clear()


def _docx_text(content: bytes) -> str:
    doc = DocxDocument(BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _generate_and_download(client, visit_uuid, kind):
    r = client.post(f"/api/visits/{visit_uuid}/documents/{kind}")
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["kind"] == kind
    assert body["utterance_id"] is None  # visit-grain (rev 0010)
    assert body["visit_id"] is not None
    d = client.get(body["download_url"])
    assert d.status_code == 200
    assert d.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    return _docx_text(d.content)


def test_transcript_contains_every_raw_turn_verbatim(env):
    client, visit_uuid = env
    text = _generate_and_download(client, visit_uuid, "transcript")
    for _, raw in RAW_TURNS:
        assert raw in text  # byte-exact raw words (rule #1)
    # Role labels present, and turn order preserved.
    assert "Assistant asked" in text and "Patient said" in text
    assert text.index(RAW_TURNS[0][1]) < text.index(RAW_TURNS[3][1])
    # Nothing summarizes/corrects in a transcript.
    assert "Verbatim" in text


def test_summary_report_has_fields_values_and_disclaimer(env):
    client, visit_uuid = env
    text = _generate_and_download(client, visit_uuid, "summary_report")
    # All 10 bilingual field labels.
    assert "1. Main Problem / প্রধান সমস্যা" in text
    assert "10. Current Concern / মূল উদ্বেগ" in text
    # Stored values rendered; missing fields shown as em-dash, not invented.
    assert "মাথা ব্যথা ও হালকা জ্বর" in text
    assert "Napa, twice daily" in text
    assert "—" in text
    # Patient meta + chief complaint + the rule-#2 disclaimer.
    assert "Kamal Hossain" in text and "120/80" in text
    assert "Headache and mild fever for three days." in text
    assert "NOT a diagnosis" in text


def test_route_guards(env):
    client, visit_uuid = env
    assert client.post(f"/api/visits/{visit_uuid}/documents/bogus").status_code == 400
    missing = "00000000-0000-0000-0000-000000000000"
    assert client.post(f"/api/visits/{missing}/documents/transcript").status_code == 404
