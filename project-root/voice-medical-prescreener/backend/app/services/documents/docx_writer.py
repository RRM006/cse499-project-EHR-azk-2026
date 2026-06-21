"""DOCX writer (python-docx).

Builds a Word document for ONE side of a session — either the verbatim RAW
transcript or the SEPARATE corrected transcript (chosen via ``kind``). Pure Python,
so it runs the same on the Windows desktop and the Arch laptop with no Word /
LibreOffice installed.

Bangla note: Bengali is a *complex script*, so we set the font on both the Latin
(``w:ascii``/``w:hAscii``) and complex-script (``w:cs``) slots of the Normal style.
That makes the saved file request a Bengali-capable font (Noto Sans Bengali, with
Nikosh / Vrinda as common fallbacks) regardless of the reader's default. We only
set the rendering font — the patient's words (rule #1) are written exactly as stored.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Pt

from backend.app.db.models import Utterance
from backend.app.services.documents.base import DocumentKind, DocumentWriter

# First name is the primary request; readers fall back to whatever Bengali font
# they have. All are common on Windows/Linux or free to install.
BENGALI_FONT = "Noto Sans Bengali"


def _apply_bengali_font(style, font_name: str = BENGALI_FONT, size_pt: int = 11) -> None:
    """Set ``font_name`` on a style for Latin AND complex (Bengali) script slots."""
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for slot in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(slot), font_name)


def _fmt_dt(value) -> str:
    return value.strftime("%Y-%m-%d %H:%M:%S UTC") if value else "—"


# Per kind: the document title and the metadata rows that belong with it. Raw and
# corrected files are independent, so each carries only its own provenance.
_TITLES: dict[str, str] = {"raw": "Transcript", "corrected": "Corrected Transcript"}


class DocxWriter(DocumentWriter):
    format = "docx"

    def render(self, utterance: Utterance, *, kind: DocumentKind) -> bytes:
        if kind not in _TITLES:
            raise ValueError(f"Unknown document kind '{kind}'. Expected 'raw' or 'corrected'.")

        doc = DocxDocument()
        _apply_bengali_font(doc.styles["Normal"])

        # --- title + separator rule (matches the requested DOCX layout) ---
        doc.add_heading(_TITLES[kind], level=1)
        rule = doc.add_paragraph("_" * 48)
        rule.runs[0].font.color.rgb = None  # keep default; just a visual divider

        # --- metadata: only what is relevant to THIS document ---
        meta = [("Session ID", str(utterance.id)), ("Source", utterance.source or "—")]
        if kind == "raw":
            meta += [
                ("STT provider", utterance.stt_provider or "—"),
                ("Recorded at", _fmt_dt(utterance.created_at)),
            ]
            body = utterance.raw_text  # verbatim, never edited (rule #1)
        else:  # corrected
            meta += [
                ("Correction", f"{utterance.correction_provider or '—'} / {utterance.correction_model or '—'}"),
                ("Corrected at", _fmt_dt(utterance.corrected_at)),
            ]
            body = utterance.corrected_text or ""

        table = doc.add_table(rows=0, cols=2)
        table.style = "Light List Accent 1"
        for label, value in meta:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value

        # --- the transcript text itself ---
        doc.add_paragraph("")  # spacer
        doc.add_paragraph(body)

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
