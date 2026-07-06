"""Visit-grain DOCX writers (rev 0010, ADR-0032): the full-visit RAW transcript
(KIOSK-4) and the staff summary report (MEDIC-7).

Different grain from ``docx_writer.DocxWriter`` (one utterance) — these render a
whole visit, so they take the visit's stored rows directly. Both are LOCAL and
deterministic: content comes only from what the pipeline already stored, so
rendering never spends quota and never invents text.

CONSTITUTION RULE #1: the transcript writer reproduces every ``raw_text`` verbatim —
no editing, paraphrasing, or re-correcting. The summary report clearly separates the
AI-derived fields from the verbatim Q&A, and always carries the M12 no-diagnosis
disclaimer (rule #2).
"""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument

from backend.app.db.models import Patient, Utterance, Visit
from backend.app.schemas.profile import SUMMARY_FIELD_KEYS
from backend.app.services.documents.docx_writer import _apply_bengali_font, _fmt_dt

#: Visit-grain values for documents.kind ("prescription" joins in the DOCTOR-6 step).
VISIT_DOCUMENT_KINDS = ("transcript", "summary_report")

# Bilingual labels for the 10 fixed fields (EN / BN, matching the kiosk's numbering).
_FIELD_LABELS: dict[str, str] = {
    "main_problem": "1. Main Problem / প্রধান সমস্যা",
    "onset_duration": "2. When Started & Duration / শুরুর সময় ও স্থায়িত্ব",
    "symptom_details": "3. Symptom Details / উপসর্গের বিস্তারিত",
    "associated_symptoms": "4. Associated Symptoms / আনুষঙ্গিক উপসর্গ",
    "medical_history": "5. Medical History / চিকিৎসা ইতিহাস",
    "current_medicines": "6. Current Medicines / চলমান ওষুধসমূহ",
    "allergies": "7. Allergies / অ্যালার্জি",
    "recent_changes_exposures": "8. Recent Changes / সাম্প্রতিক পরিবর্তন",
    "treatments_tried": "9. Treatments Tried / গৃহীত ব্যবস্থা",
    "current_concern": "10. Current Concern / মূল উদ্বেগ",
}

_SPEAKER_LABELS = {"system": "Assistant asked", "patient": "Patient said"}


def _new_doc() -> DocxDocument:
    doc = DocxDocument()
    _apply_bengali_font(doc.styles["Normal"])
    return doc


def _meta_table(doc: DocxDocument, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _patient_meta(patient: Patient | None) -> list[tuple[str, str]]:
    if patient is None:
        return [("Patient", "—")]
    rows = [
        ("Patient", patient.display_name or "—"),
        ("Phone", patient.external_ref or "—"),
    ]
    if patient.sex:
        rows.append(("Sex", patient.sex))
    if patient.birth_year:
        rows.append(("Birth year", str(patient.birth_year)))
    if patient.weight_kg is not None:
        rows.append(("Weight", f"{patient.weight_kg:g} kg"))
    if patient.bp:
        rows.append(("Blood pressure", patient.bp))
    return rows


def _to_bytes(doc: DocxDocument) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_visit_transcript(
    visit: Visit, patient: Patient | None, utterances: list[Utterance]
) -> bytes:
    """The whole conversation, verbatim, in turn order (KIOSK-4)."""
    doc = _new_doc()
    doc.add_heading("Visit Transcript — Raw (Verbatim)", level=1)
    doc.add_paragraph(
        "This document reproduces the captured speech exactly as recognized. "
        "Nothing has been edited, corrected, or summarized."
    )
    _meta_table(
        doc,
        [
            ("Visit ID", visit.uuid),
            *_patient_meta(patient),
            ("Started at", _fmt_dt(visit.started_at)),
            ("Turns", str(len(utterances))),
        ],
    )
    doc.add_paragraph("")
    for u in utterances:
        speaker = doc.add_paragraph()
        speaker.add_run(
            _SPEAKER_LABELS.get(u.role, u.role) + ":"
        ).bold = True
        doc.add_paragraph(u.raw_text)  # exact words, never edited (rule #1)
    return _to_bytes(doc)


def _field_value(field) -> str:
    """One summary field's display text; tolerates the bilingual shape
    (value/value_en/value_bn, Session 9) and missing/legacy shapes."""
    if isinstance(field, dict):
        for slot in ("value", "value_en", "value_bn"):
            text = str(field.get(slot) or "").strip()
            if text:
                return text
    return "—"


def render_visit_summary_report(
    visit: Visit, patient: Patient | None, sections: dict
) -> bytes:
    """The staff-facing pre-screening summary (MEDIC-7), from the stored M12
    ``Report.sections`` — AI-derived content is labeled as such; the follow-up
    answers stay verbatim; the no-diagnosis disclaimer is always present."""
    doc = _new_doc()
    doc.add_heading("Pre-Screening Summary Report", level=1)
    _meta_table(
        doc,
        [
            ("Visit ID", visit.uuid),
            *_patient_meta(patient),
            ("Started at", _fmt_dt(visit.started_at)),
        ],
    )

    risk = sections.get("risk") or {}
    tier = risk.get("tier")
    doc.add_heading("Risk Assessment", level=2)
    doc.add_paragraph(f"Tier: {tier or 'not assessed'}")
    red_flags = sections.get("red_flags") or []
    if red_flags:
        doc.add_paragraph("RED FLAGS (rule-forced, verify immediately):")
        for flag in red_flags:
            doc.add_paragraph(f"• {flag}")
    if sections.get("xai"):
        doc.add_paragraph(f"Reasoning: {sections['xai']}")

    if sections.get("chief_complaint"):
        doc.add_heading("Chief Complaint (AI summary)", level=2)
        doc.add_paragraph(sections["chief_complaint"])

    # C1 (ADR-0036): the labeled suggestion — never rendered without its disclaimer.
    suggestion = sections.get("suggested_condition") or {}
    condition_en = str(suggestion.get("condition_en") or suggestion.get("condition") or "").strip()
    condition_bn = str(suggestion.get("condition_bn") or "").strip()
    condition = condition_en or condition_bn
    if condition:
        doc.add_heading(
            "Possible Condition (AI Suggestion – Not a Diagnosis) / "
            "সম্ভাব্য অবস্থা (এআই পরামর্শ – রোগনির্ণয় নয়)",
            level=2,
        )
        if condition_bn and condition_bn != condition:
            condition += f" / {condition_bn}"
        para = doc.add_paragraph()
        para.add_run(condition).bold = True
        reasoning = str(
            suggestion.get("reasoning_en") or suggestion.get("reasoning_bn") or ""
        ).strip()
        if reasoning:
            doc.add_paragraph(f"Reasoning: {reasoning}")
        c1_disclaimer = doc.add_paragraph()
        c1_disclaimer.add_run(
            str(suggestion.get("disclaimer") or "") + " "
            + str(suggestion.get("disclaimer_bn") or "")
        ).italic = True

    doc.add_heading("Pre-Screening Summary (10 fields, AI-extracted)", level=2)
    fields = sections.get("summary_fields") or {}
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for key in SUMMARY_FIELD_KEYS:
        cells = table.add_row().cells
        cells[0].text = _FIELD_LABELS[key]
        cells[1].text = _field_value(fields.get(key))

    qa = sections.get("followup_qa") or []
    if qa:
        doc.add_heading("Follow-up Questions & Answers (answers verbatim)", level=2)
        for item in qa:
            q = doc.add_paragraph()
            q.add_run("Q: ").bold = True
            q.add_run(str(item.get("question") or ""))
            a = doc.add_paragraph()
            a.add_run("A: ").bold = True
            a.add_run(str(item.get("answer_raw") or "(no answer recorded)"))

    doc.add_paragraph("")
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(str(sections.get("disclaimer") or "")).italic = True
    return _to_bytes(doc)


def _rx(payload: dict, key: str) -> str:
    """One prescription payload field as a trimmed string ('' if missing)."""
    return str(payload.get(key) or "").strip()


def render_prescription(payload: dict) -> bytes:
    """DOCTOR-6: a doctor-authored prescription rendered from the form ``payload``.

    LOCAL and self-contained — the payload IS the whole document. The Diagnosis is
    printed exactly as the doctor typed it (``payload['diagnosis']``); this writer
    never reads the AI suggested condition, so a diagnosis can never be AI-filled
    (constitution rule #2). Raw utterances are not touched (rule #1).
    """
    doc = _new_doc()

    clinic = payload.get("clinic") or {}
    doctor = payload.get("doctor") or {}
    patient = payload.get("patient") or {}

    # --- Letterhead ---
    doc.add_heading(str(clinic.get("name") or "Clinic"), level=1)
    if clinic.get("address"):
        doc.add_paragraph(str(clinic["address"]))
    doc_line = " · ".join(
        part for part in [
            str(doctor.get("name") or ""),
            str(doctor.get("qualification") or ""),
            str(doctor.get("specialization") or ""),
            (f"Reg: {doctor['registration_no']}" if doctor.get("registration_no") else ""),
        ] if part
    )
    if doc_line:
        head = doc.add_paragraph()
        head.add_run(doc_line).bold = True
    doc.add_paragraph(f"Date: {_rx(payload, 'date') or '—'}")

    # --- Patient ---
    patient_bits = " · ".join(
        part for part in [
            str(patient.get("name") or ""),
            str(patient.get("phone") or ""),
            (f"Age {patient['age']}" if patient.get("age") else ""),
            str(patient.get("sex") or ""),
            (f"{patient['weight_kg']} kg" if patient.get("weight_kg") else ""),
            (f"BP {patient['bp']}" if patient.get("bp") else ""),
        ] if part
    )
    _meta_table(doc, [("Patient", patient_bits or "—")])

    if _rx(payload, "symptoms"):
        doc.add_heading("Symptoms / Chief Complaints", level=2)
        doc.add_paragraph(_rx(payload, "symptoms"))

    # Diagnosis — printed verbatim from the payload; NEVER AI-filled (rule #2).
    doc.add_heading("Diagnosis", level=2)
    doc.add_paragraph(_rx(payload, "diagnosis") or "—")

    medicines = [m for m in (payload.get("medicines") or []) if str(m.get("name") or "").strip()]
    if medicines:
        doc.add_heading("Rx — Medicines", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light List Accent 1"
        header = table.rows[0].cells
        for i, label in enumerate(("Name", "Strength", "Dosage", "Timing", "Duration")):
            header[i].text = label
        for m in medicines:
            cells = table.add_row().cells
            cells[0].text = str(m.get("name") or "")
            cells[1].text = str(m.get("strength") or "")
            cells[2].text = str(m.get("dosage") or "")
            cells[3].text = str(m.get("timing") or "")
            cells[4].text = str(m.get("duration") or "")

    if _rx(payload, "advice"):
        doc.add_heading("Advice / Lifestyle", level=2)
        doc.add_paragraph(_rx(payload, "advice"))
    if _rx(payload, "tests"):
        doc.add_heading("Required Tests", level=2)
        doc.add_paragraph(_rx(payload, "tests"))
    if _rx(payload, "followup_date"):
        doc.add_paragraph(f"Follow-up: {_rx(payload, 'followup_date')}")

    doc.add_paragraph("")
    sig = doc.add_paragraph()
    sig.add_run(f"{doctor.get('name') or ''}\nSignature & Stamp").italic = True
    return _to_bytes(doc)
